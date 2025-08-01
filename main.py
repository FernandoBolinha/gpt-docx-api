from fastapi import FastAPI, Request
from pydantic import BaseModel
from fastapi.responses import FileResponse
from docx import Document
import uuid
import os

app = FastAPI()

# Pasta onde os arquivos serão salvos
OUTPUT_DIR = "docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

class WordRequest(BaseModel):
    titulo: str
    conteudo: str

@app.post("/gerar-docx")
def gerar_docx(data: WordRequest):
    file_id = str(uuid.uuid4())
    filename = f"{file_id}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    doc.add_heading(data.titulo, 0)
    doc.add_paragraph(data.conteudo)
    doc.save(filepath)

    return {
        "download_url": f"/download/{filename}"
    }

@app.get("/download/{filename}")
def download(filename: str):
    filepath = os.path.join(OUTPUT_DIR, filename)
    return FileResponse(filepath, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
